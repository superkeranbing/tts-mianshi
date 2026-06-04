import json, re, random
from typing import Optional

# Simulated ASR engine
class ASREngine:
    """模拟 ASR 引擎，生产环境接入 FunASR/Whisper"""
    def __init__(self):
        self.backend = "funasr_sim"

    async def transcribe_file(self, audio_path: str) -> list[dict]:
        """离线转写音频文件，返回带时间戳的段落列表。生产环境调用 FunASR。"""
        # Return mock transcription
        return [
            {"speaker": "面试官", "text": "请简单介绍一下你自己。", "start": 0.0, "end": 3.5},
            {"speaker": "候选人", "text": "我叫张三，毕业于XX大学，有5年前端开发经验。", "start": 4.0, "end": 18.0},
            {"speaker": "面试官", "text": "React的虚拟DOM原理是什么？", "start": 20.0, "end": 25.0},
            {"speaker": "候选人", "text": "虚拟DOM是React用JS对象模拟真实DOM的优化机制...", "start": 26.0, "end": 52.0},
            {"speaker": "面试官", "text": "你在项目中遇到的最大技术挑战？", "start": 54.0, "end": 58.0},
            {"speaker": "候选人", "text": "大数据量列表渲染性能问题，我们采用了虚拟滚动方案...", "start": 59.0, "end": 85.0},
            {"speaker": "面试官", "text": "你的职业规划是什么？", "start": 87.0, "end": 90.0},
            {"speaker": "候选人", "text": "计划3年内成为全栈架构师，深耕前端并拓展后端能力。", "start": 91.0, "end": 110.0},
        ]

    async def stream_transcribe(self, audio_chunk: bytes, previous_text: str = "") -> dict:
        """流式转写，返回部分识别结果。生产环境使用 FunASR 流式接口。"""
        return {"text": previous_text + "…", "is_final": False}

    async def detect_language(self, audio_path: str) -> str:
        """检测音频语种"""
        return "zh"

    async def vad_segment(self, audio_path: str) -> list[dict]:
        """语音活动检测，切分音频"""
        return [{"start": 0, "end": 130}]


class DiarizationEngine:
    """说话人分离引擎"""
    async def separate_speakers(self, audio_path: str, segments: list[dict]) -> list[dict]:
        """为每个转写段落标注说话人。生产环境使用 CAM++。"""
        speakers = ["面试官", "候选人"]
        for i, seg in enumerate(segments):
            seg["speaker"] = speakers[i % 2]
        return segments


class SummaryService:
    """智能纪要生成"""
    async def generate_summary(self, transcripts: list[dict]) -> str:
        return "本次面试持续约20分钟，面试官主要考察了候选人的前端技术基础、项目经验和职业规划。候选人表现良好，技术基础扎实，但在系统设计方面仍有提升空间。"


class ExportService:
    """导出服务"""
    async def export_txt(self, title: str, transcripts: list[dict]) -> str:
        lines = [f"标题: {title}\n"]
        for t in transcripts:
            ts = f"[{int(t.get('start',0)//60):02d}:{int(t.get('start',0)%60):02d}]"
            lines.append(f"{ts} {t.get('speaker','未知')}: {t.get('content','')}")
        return "\n".join(lines)

    async def export_srt(self, transcripts: list[dict]) -> str:
        lines = []
        for i, t in enumerate(transcripts, 1):
            speaker = t.get("speaker", "未知")
            content = t.get("content", "")
            start = _format_srt(t.get("start", 0))
            end = _format_srt(t.get("end", 0))
            lines.append(f"{i}\n{start} --> {end}\n{speaker}: {content}\n")
        return "\n".join(lines)

    async def export_docx(self, title: str, transcripts: list[dict]) -> bytes:
        try:
            from docx import Document
            doc = Document()
            doc.add_heading(title, 0)
            for t in transcripts:
                doc.add_paragraph(f"{t.get('speaker','未知')}: {t.get('content','')}")
            import io
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except ImportError:
            return b"python-docx not installed"

def _format_srt(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds % 1) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"
