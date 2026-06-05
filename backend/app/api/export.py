from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import select
import io, os, json
from app.core.database import get_db
from app.models.recording import Recording
from app.models.transcript import Transcript
from app.models.interview import InterviewReport, QAPair, KnowledgePoint
from app.config import get_settings

router = APIRouter(prefix="/api/export", tags=["Export"])
settings = get_settings()


def _srt(sec):
    m = int(sec // 60)
    s = int(sec % 60)
    ms = int(sec % 1 * 1000)
    return f"{m:02d}:{s:02d}:{ms:03d}"


def _fmt_mmss(sec):
    return f"{int(sec//60):02d}:{int(sec%60):02d}"


def _load_transcripts(recording_id: str, db: Session):
    recording = db.execute(select(Recording).where(Recording.id == recording_id)).scalar_one_or_none()
    if not recording:
        raise HTTPException(404, "Recording not found")
    tl = db.execute(
        select(Transcript)
        .where(Transcript.recording_id == recording_id)
        .order_by(Transcript.start_time)
    ).scalars().all()
    return recording, tl


@router.get("/{recording_id}/txt")
async def export_txt(recording_id: str, db: Session = Depends(get_db)):
    recording, tl = _load_transcripts(recording_id, db)
    header = f"Title: {recording.title}" + chr(10)
    lines = [header]
    for t in tl:
        sp = t.speaker_name or t.speaker
        ts = _fmt_mmss(t.start_time)
        lines.append(f"{ts} {sp}: " + t.content + chr(10))
    content = chr(10).join(lines)
    return StreamingResponse(
        io.BytesIO(content.encode("utf-8")),
        media_type="text/plain",
        headers={"Content-Disposition": f"attachment; filename={recording.title}.txt"},
    )


@router.get("/{recording_id}/srt")
async def export_srt(recording_id: str, db: Session = Depends(get_db)):
    recording, tl = _load_transcripts(recording_id, db)
    lines = []
    for i, t in enumerate(tl, 1):
        sp = t.speaker_name or t.speaker
        srt_line = f"{i}" + chr(10) + _srt(t.start_time) + " --> " + _srt(t.end_time) + chr(10) + sp + ": " + t.content + chr(10)
        lines.append(srt_line)
    content = chr(10).join(lines)
    return StreamingResponse(
        io.BytesIO(content.encode("utf-8")),
        media_type="text/plain",
        headers={"Content-Disposition": f"attachment; filename={recording.title}.srt"},
    )


@router.get("/{recording_id}/docx")
async def export_docx(recording_id: str, db: Session = Depends(get_db)):
    from docx import Document
    from docx.shared import Pt, RGBColor
    recording, tl = _load_transcripts(recording_id, db)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    doc.add_heading(recording.title, level=1)
    dur_text = f"Duration: {recording.audio_duration:.1f}s" if recording.audio_duration else ""
    doc.add_paragraph(dur_text)
    for t in tl:
        sp = t.speaker_name or t.speaker
        ts = _fmt_mmss(t.start_time)
        p = doc.add_paragraph()
        runner = p.add_run(f"[{ts}] {sp}")
        runner.bold = True
        runner.font.size = Pt(10)
        runner.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
        content_run = p.add_run(chr(10) + t.content)
        content_run.font.size = Pt(10)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={recording.title}.docx"},
    )


@router.get("/{recording_id}/pdf")
async def export_pdf(recording_id: str, db: Session = Depends(get_db)):
    from fpdf import FPDF
    recording, tl = _load_transcripts(recording_id, db)
    pdf = FPDF()
    pdf.add_page()
    font_path = os.path.join(os.path.dirname(__file__), "..", "utils", "NotoSansSC-Regular.ttf")
    if os.path.exists(font_path):
        pdf.add_font("notosans", "", font_path, uni=True)
        family = "notosans"
    else:
        family = "Helvetica"
    pdf.set_font(family, "", 16)
    pdf.set_text_color(50, 50, 50)
    pdf.cell(0, 12, recording.title, new_x="LMARGIN", new_y="NEXT")
    if recording.audio_duration:
        pdf.set_font(family, "", 9)
        pdf.set_text_color(130, 130, 130)
        pdf.cell(0, 8, f"Duration: {recording.audio_duration:.1f}s", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    for t in tl:
        sp = t.speaker_name or t.speaker
        ts = _fmt_mmss(t.start_time)
        pdf.set_font(family, "", 10)
        pdf.set_text_color(16, 185, 129)
        pdf.cell(0, 7, f"[{ts}] {sp}", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(180, 180, 180)
        pdf.set_font(family, "", 10)
        pdf.multi_cell(0, 6, t.content)
        pdf.ln(2)
    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={recording.title}.pdf"},
    )


@router.get("/report/{report_id}/pdf")
async def export_report_pdf(report_id: str, db: Session = Depends(get_db)):
    from fpdf import FPDF
    report = db.execute(select(InterviewReport).where(InterviewReport.id == report_id)).scalar_one_or_none()
    if not report:
        raise HTTPException(404, "Report not found")
    qa_pairs = db.execute(select(QAPair).where(QAPair.report_id == report_id)).scalars().all()
    kps = db.execute(select(KnowledgePoint).where(KnowledgePoint.report_id == report_id)).scalars().all()
    strengths = json.loads(report.strengths_json) if report.strengths_json else []
    weaknesses = json.loads(report.weaknesses_json) if report.weaknesses_json else []
    plan = json.loads(report.improvement_plan_json) if report.improvement_plan_json else []
    font_path = os.path.join(os.path.dirname(__file__), "..", "utils", "NotoSansSC-Regular.ttf")
    pdf = FPDF()
    pdf.add_page()
    if os.path.exists(font_path):
        pdf.add_font("notosans", "", font_path, uni=True)
        family = "notosans"
    else:
        family = "Helvetica"
    pdf.set_font(family, "", 20)
    pdf.set_text_color(50, 50, 50)
    pdf.cell(0, 15, "Interview Report", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(family, "", 10)
    pdf.set_text_color(130, 130, 130)
    pdf.cell(0, 8, f"Score: {report.overall_score}/100" if report.overall_score else "", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    if report.summary:
        pdf.set_font(family, "", 11)
        pdf.set_text_color(80, 80, 80)
        pdf.multi_cell(0, 7, f"Summary: {report.summary}")
        pdf.ln(4)
    if strengths:
        pdf.set_font(family, "", 12)
        pdf.set_text_color(16, 185, 129)
        pdf.cell(0, 10, "Strengths", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(100, 100, 100)
        pdf.set_font(family, "", 10)
        for s in strengths:
            pdf.cell(0, 7, f"  + {s}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)
    if weaknesses:
        pdf.set_font(family, "", 12)
        pdf.set_text_color(250, 204, 21)
        pdf.cell(0, 10, "Improvements", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(100, 100, 100)
        pdf.set_font(family, "", 10)
        for w in weaknesses:
            pdf.cell(0, 7, f"  - {w}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)
    if qa_pairs:
        pdf.set_font(family, "", 14)
        pdf.set_text_color(60, 60, 60)
        pdf.cell(0, 12, "Q&A Analysis", new_x="LMARGIN", new_y="NEXT")
        for i, qa in enumerate(qa_pairs, 1):
            if pdf.get_y() > 250:
                pdf.add_page()
            pdf.set_font(family, "", 11)
            pdf.set_text_color(16, 185, 129)
            pdf.cell(0, 8, f"Q{i}: {qa.question}", new_x="LMARGIN", new_y="NEXT")
            if qa.your_answer:
                pdf.set_text_color(140, 140, 140)
                pdf.set_font(family, "", 9)
                pdf.multi_cell(0, 5, f"Your answer: {qa.your_answer[:200]}")
            if qa.best_answer:
                pdf.set_text_color(16, 185, 129)
                pdf.set_font(family, "", 9)
                pdf.multi_cell(0, 5, f"Best answer: {qa.best_answer[:200]}")
            pdf.ln(3)
    if kps:
        if pdf.get_y() > 200:
            pdf.add_page()
        pdf.set_font(family, "", 14)
        pdf.set_text_color(60, 60, 60)
        pdf.cell(0, 12, "Knowledge Points", new_x="LMARGIN", new_y="NEXT")
        for kp in kps:
            concepts = json.loads(kp.key_concepts_json) if kp.key_concepts_json else []
            if pdf.get_y() > 240:
                pdf.add_page()
            pdf.set_font(family, "", 11)
            pdf.set_text_color(16, 185, 129)
            pdf.cell(0, 8, kp.title, new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(100, 100, 100)
            pdf.set_font(family, "", 9)
            pdf.multi_cell(0, 5, kp.content[:300])
            if concepts:
                pdf.cell(0, 6, "Keywords: " + ", ".join(concepts), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(3)
    if plan:
        if pdf.get_y() > 240:
            pdf.add_page()
        pdf.set_font(family, "", 12)
        pdf.set_text_color(60, 60, 60)
        pdf.cell(0, 10, "Improvement Plan", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(100, 100, 100)
        pdf.set_font(family, "", 10)
        for p in plan:
            pdf.cell(0, 7, f"Week " + str(p.get("week","")) + ": " + p.get("focus",""), new_x="LMARGIN", new_y="NEXT")
    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=interview_report_{report_id}.pdf"},
    )